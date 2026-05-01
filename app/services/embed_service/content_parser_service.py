"""File content parsing utilities for RAG ingestion / chunking.

This module takes raw bytes plus lightweight hints (filename / content-type)
and returns cleaned, plain text suitable for downstream chunking.
"""

from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from typing import Literal

from app.services.string_helper_service import collapse_whitespace


class ContentParseError(RuntimeError):
    pass


class UnsupportedFileTypeError(ContentParseError):
    pass


class MissingDependencyError(ContentParseError):
    pass


ParsedKind = Literal["pdf", "doc", "docx", "pptx", "xlsx", "image", "text"]


@dataclass(frozen=True)
class ParsedContent:
    text: str
    kind: ParsedKind
    detected_mime: str | None = None
    filename: str | None = None


def parse_file_bytes(
    content: bytes,
    *,
    filename: str | None = None,
    content_type: str | None = None,
) -> ParsedContent:
    """Parse bytes into cleaned text.

    Args:
        content: Raw file bytes.
        filename: Optional filename (used for extension-based detection).
        content_type: Optional MIME type hint (e.g. "application/pdf").
    """
    if not content:
        return ParsedContent(text="", kind="text", detected_mime=content_type, filename=filename)

    kind, detected_mime = _detect_kind(content, filename=filename, content_type=content_type)

    if kind == "pdf":
        raw = _parse_pdf(content)
    elif kind == "docx":
        raw = _parse_docx(content)
    elif kind == "doc":
        raw = _parse_doc(content)
    elif kind == "pptx":
        raw = _parse_office_via_textract(content, extension="pptx")
    elif kind == "xlsx":
        raw = _parse_office_via_textract(content, extension="xlsx")
    elif kind == "image":
        raw = _parse_image_ocr(content)
    elif kind == "text":
        raw = _parse_text(content, filename=filename, content_type=content_type)
    else:  # pragma: no cover
        raise UnsupportedFileTypeError(f"Unsupported kind: {kind}")

    return ParsedContent(
        text=_clean_text(raw),
        kind=kind,
        detected_mime=detected_mime or content_type,
        filename=filename,
    )


def _detect_kind(
    content: bytes,
    *,
    filename: str | None,
    content_type: str | None,
) -> tuple[ParsedKind, str | None]:
    ct = (content_type or "").split(";")[0].strip().lower()

    if ct in {"application/pdf"}:
        return "pdf", ct
    if ct in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    }:
        # msword can be doc or docx; fall through to sniffing.
        pass
    if ct.startswith("image/"):
        return "image", ct
    if ct.startswith("text/"):
        return "text", ct

    ext = (filename or "").lower().rsplit(".", 1)[-1] if filename and "." in filename else ""
    if ext == "pdf":
        return "pdf", "application/pdf"
    if ext == "docx":
        return (
            "docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    if ext == "doc":
        return "doc", "application/msword"
    if ext == "pptx":
        return "pptx", "application/vnd.openxmlformats-officedocument.presentationml.presentation"
    if ext == "xlsx":
        return "xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    if ext in {"txt", "md", "rtf", "csv", "json", "xml", "html", "htm"}:
        return "text", "text/plain"
    if ext in {"png", "jpg", "jpeg", "gif", "tif", "tiff", "bmp", "webp"}:
        return "image", f"image/{'jpeg' if ext in {'jpg', 'jpeg'} else ext}"

    # Magic bytes sniffing (kept intentionally lightweight).
    if content.startswith(b"%PDF-"):
        return "pdf", "application/pdf"
    if content[:8].startswith(b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1"):
        return "doc", "application/msword"
    if content[:2] == b"PK":
        # Could be many ZIP-based formats; treat as docx if we can't tell otherwise.
        return (
            "docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    if _looks_like_image(content):
        return "image", None

    return "text", ct or None


def _looks_like_image(content: bytes) -> bool:
    # PNG
    if content.startswith(b"\x89PNG\r\n\x1a\n"):
        return True
    # JPEG
    if content.startswith(b"\xff\xd8\xff"):
        return True
    # GIF
    if content.startswith(b"GIF87a") or content.startswith(b"GIF89a"):
        return True
    # TIFF
    if content.startswith(b"II*\x00") or content.startswith(b"MM\x00*"):
        return True
    # BMP
    if content.startswith(b"BM"):
        return True
    # WEBP (RIFF....WEBP)
    if content.startswith(b"RIFF") and b"WEBP" in content[:16]:
        return True
    return False


def _parse_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception as exc:  # pragma: no cover
        raise MissingDependencyError(
            "PDF parsing requires `pypdf`. Add it to requirements and install dependencies."
        ) from exc

    reader = PdfReader(BytesIO(content))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return "\n\n".join(parts)


def _parse_docx(content: bytes) -> str:
    try:
        import docx  # type: ignore
    except Exception as exc:  # pragma: no cover
        raise MissingDependencyError(
            "DOCX parsing requires `python-docx`. Add it to requirements and install dependencies."
        ) from exc

    document = docx.Document(BytesIO(content))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text and paragraph.text.strip():
            parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join((cell.text or "").strip() for cell in row.cells).strip()
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def _parse_doc(content: bytes) -> str:
    # Best-effort: `textract` can handle .doc, but may require system deps.
    try:
        import textract  # type: ignore
    except Exception as exc:  # pragma: no cover
        raise MissingDependencyError(
            "DOC parsing requires `textract` (and may require system packages like antiword)."
        ) from exc

    import os
    import tempfile

    tmp_path: str | None = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as tmp:
            tmp.write(content)
            tmp.flush()
            tmp_path = tmp.name
        extracted = textract.process(tmp_path)
        return extracted.decode("utf-8", errors="replace")
    finally:
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass


def _parse_office_via_textract(content: bytes, *, extension: str) -> str:
    try:
        import textract  # type: ignore
    except Exception as exc:  # pragma: no cover
        raise MissingDependencyError(
            f"{extension.upper()} parsing requires `textract` (and may require system packages)."
        ) from exc

    import os
    import tempfile

    tmp_path: str | None = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{extension}") as tmp:
            tmp.write(content)
            tmp.flush()
            tmp_path = tmp.name
        extracted = textract.process(tmp_path)
        return extracted.decode("utf-8", errors="replace")
    finally:
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass


def _parse_image_ocr(content: bytes) -> str:
    try:
        from PIL import Image  # type: ignore
    except Exception as exc:  # pragma: no cover
        raise MissingDependencyError(
            "Image parsing requires `Pillow`. Add it to requirements and install dependencies."
        ) from exc

    try:
        import pytesseract  # type: ignore
    except Exception as exc:  # pragma: no cover
        raise MissingDependencyError(
            "Image OCR requires `pytesseract` and the `tesseract` binary available in the runtime."
        ) from exc

    img = Image.open(BytesIO(content))
    return pytesseract.image_to_string(img)


def _parse_text(content: bytes, *, filename: str | None, content_type: str | None) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            text = content.decode(encoding)
            return _maybe_strip_html(text, filename=filename, content_type=content_type)
        except Exception:
            continue
    text = content.decode("utf-8", errors="replace")
    return _maybe_strip_html(text, filename=filename, content_type=content_type)


def _maybe_strip_html(text: str, *, filename: str | None, content_type: str | None) -> str:
    ct = (content_type or "").split(";")[0].strip().lower()
    ext = (filename or "").lower().rsplit(".", 1)[-1] if filename and "." in filename else ""
    if ct in {"text/html", "application/xhtml+xml"} or ext in {"html", "htm"}:
        return _strip_html(text)
    return text


def _strip_html(html: str) -> str:
    import re

    html = re.sub(r"<script[\s\S]*?</script>", " ", html, flags=re.IGNORECASE)
    html = re.sub(r"<style[\s\S]*?</style>", " ", html, flags=re.IGNORECASE)
    html = re.sub(r"<[^>]+>", " ", html)
    html = (
        html.replace("&nbsp;", " ")
        .replace("&amp;", "&")
        .replace("&lt;", "<")
        .replace("&gt;", ">")
        .replace("&quot;", '"')
        .replace("&#39;", "'")
    )
    html = re.sub(r"\s+", " ", html)
    return html.strip()


def _clean_text(text: str) -> str:
    if not text:
        return ""
    text = text.replace("\x00", "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Normalize whitespace inside lines while preserving line breaks for chunking.
    lines = []
    for line in text.split("\n"):
        cleaned = collapse_whitespace(line)
        lines.append(cleaned)
    # Collapse runs of blank lines down to a single blank line.
    out_lines: list[str] = []
    blank_run = 0
    for line in lines:
        if not line:
            blank_run += 1
            if blank_run <= 1:
                out_lines.append("")
            continue
        blank_run = 0
        out_lines.append(line)
    return "\n".join(out_lines).strip()
